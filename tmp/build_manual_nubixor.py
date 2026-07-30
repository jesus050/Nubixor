from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/jesusdiaz/Documents/Nubixor")
OUTPUT = ROOT / "docs" / "Manual_Funcional_Nubixor.docx"
LOGO = ROOT / "public" / "assets" / "brand" / "nubixor-logo.png"

NAVY = "081447"
BLUE = "1E42BB"
CYAN = "21C9E8"
VIOLET = "8B4DFF"
MAGENTA = "DC3AD7"
INK = "101B4F"
MUTED = "687087"
PALE = "F3F6FF"
PALE_CYAN = "E7FAFD"
PALE_VIOLET = "F2ECFF"
LINE = "DDE3F3"
WHITE = "FFFFFF"
GREEN = "137C68"
AMBER = "9A6200"
RED = "A52A5D"

BODY_FONT = "Calibri"
DISPLAY_FONT = "Avenir Next Condensed"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run_font(run, name=BODY_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_text(cell, text, *, bold=False, color=INK, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, after=0, line=1.15)
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)


def add_field(paragraph, field_code):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, display, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_custom_numbering(document):
    numbering = document.part.numbering_part.element

    def abstract_num(abstract_id, fmt, text, marker_font=None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if marker_font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), marker_font)
            fonts.set(qn("w:hAnsi"), marker_font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

    def concrete_num(num_id, abstract_id):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract = OxmlElement("w:abstractNumId")
        abstract.set(qn("w:val"), str(abstract_id))
        num.append(abstract)
        numbering.append(num)

    abstract_num(20, "bullet", "•", BODY_FONT)
    abstract_num(21, "decimal", "%1.")
    concrete_num(20, 20)
    concrete_num(21, 21)
    return 20, 21


def set_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(.492)
section.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = BODY_FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
]:
    style = styles[style_name]
    style.font.name = DISPLAY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), DISPLAY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), DISPLAY_FONT)
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

bullet_num_id, decimal_num_id = add_custom_numbering(doc)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_spacing(hp, after=0, line=1)
hr = hp.add_run("NUBIXOR  /  MANUAL FUNCIONAL")
set_run_font(hr, DISPLAY_FONT, 8.5, NAVY, True)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_paragraph_spacing(fp, after=0, line=1)
fr = fp.add_run("Nubixor · Julio de 2026   |   ")
set_run_font(fr, BODY_FONT, 8.5, MUTED)
add_field(fp, "PAGE")


def add_page_break():
    doc.add_page_break()


def add_kicker(text, color=VIOLET, align=WD_ALIGN_PARAGRAPH.LEFT, after=7):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=1)
    run = p.add_run(text.upper())
    set_run_font(run, BODY_FONT, 9, color, True)
    run.font.letter_spacing = Pt(1.2)
    return p


def add_title(text, size=28, color=NAVY, align=WD_ALIGN_PARAGRAPH.LEFT, after=8):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=.92)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, DISPLAY_FONT, size, color, True)
    return p


def add_subtitle(text, size=13, color=MUTED, align=WD_ALIGN_PARAGRAPH.LEFT, after=14):
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_spacing(p, after=after, line=1.2)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, size, color)
    return p


def add_body(text, *, bold_lead=None, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after, line=1.25)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        set_run_font(run, BODY_FONT, 11, INK, True)
        run = p.add_run(text[len(bold_lead):])
        set_run_font(run, BODY_FONT, 11, INK)
    else:
        run = p.add_run(text)
        set_run_font(run, BODY_FONT, 11, INK)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph()
        set_num(p, bullet_num_id)
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run, BODY_FONT, 10.5, INK)


def add_steps(items):
    for item in items:
        p = doc.add_paragraph()
        set_num(p, decimal_num_id)
        set_paragraph_spacing(p, after=4, line=1.25)
        run = p.add_run(item)
        set_run_font(run, BODY_FONT, 10.5, INK)


def add_callout(label, text, tone="info"):
    colors = {
        "info": (PALE_CYAN, NAVY, CYAN),
        "warning": ("FFF6E5", AMBER, "F0B429"),
        "risk": ("FCEBF2", RED, MAGENTA),
        "success": ("EAF8F4", GREEN, "2FB498"),
        "violet": (PALE_VIOLET, NAVY, VIOLET),
    }
    fill, text_color, accent = colors[tone]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 180, 130, 180)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "24")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.18)
    r1 = p.add_run(f"{label.upper()}  ")
    set_run_font(r1, BODY_FONT, 9.5, text_color, True)
    r2 = p.add_run(text)
    set_run_font(r2, BODY_FONT, 9.5, text_color)
    after = doc.add_paragraph()
    set_paragraph_spacing(after, after=2, line=1)


def add_table(headers, rows, widths, *, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header_text, bold=True, color=WHITE, size=9)
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, PALE)
        for idx, value in enumerate(row_data):
            set_cell_text(row.cells[idx], value, size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    set_paragraph_spacing(after, before=0, after=3, line=1)
    return table


def add_flow(steps):
    widths = [CONTENT_WIDTH_DXA // len(steps)] * len(steps)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=len(steps))
    set_table_geometry(table, widths)
    for idx, step in enumerate(steps):
        cell = table.cell(0, idx)
        set_cell_shading(cell, NAVY if idx % 2 == 0 else BLUE)
        set_cell_margins(cell, 130, 90, 130, 90)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0, line=1.1)
        number = p.add_run(f"{idx + 1:02d}\n")
        set_run_font(number, BODY_FONT, 8, CYAN, True)
        text = p.add_run(step)
        set_run_font(text, BODY_FONT, 9, WHITE, True)
    doc.add_paragraph()


# Portada editorial
doc.add_paragraph().paragraph_format.space_after = Pt(42)
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(4.65))
    logo_shape._inline.docPr.set("descr", "Logotipo de Nubixor")
    logo_shape._inline.docPr.set("title", "Nubixor")
    set_paragraph_spacing(p, after=30, line=1)
add_kicker("Manual funcional y operativo", CYAN, WD_ALIGN_PARAGRAPH.CENTER, 14)
add_title("Cómo funciona Nubixor", 34, NAVY, WD_ALIGN_PARAGRAPH.CENTER, 10)
add_subtitle(
    "Guía completa de módulos, usuarios, procesos empresariales, controles y preparación para operación real.",
    14,
    MUTED,
    WD_ALIGN_PARAGRAPH.CENTER,
    30,
)
add_callout(
    "Versión del documento",
    "Edición 1.0 · Estado del producto revisado al 28 de julio de 2026.",
    "violet",
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_spacing(p, before=34, after=0, line=1)
run = p.add_run("ERP MULTIEMPRESA · CAJA · INVENTARIO · FINANZAS · AUDITORÍA")
set_run_font(run, BODY_FONT, 9, NAVY, True)

add_page_break()

# Índice y lectura rápida
add_kicker("Guía de navegación")
add_title("Contenido del manual", 28)
add_subtitle("El documento sigue el recorrido real de una operación, desde la configuración inicial hasta el cierre y la auditoría.")
contents = [
    ("01", "Qué es Nubixor y cómo está organizado"),
    ("02", "Acceso, usuarios, roles y aislamiento multiempresa"),
    ("03", "Configuración inicial de una empresa"),
    ("04", "Dashboard ejecutivo"),
    ("05", "Catálogo de productos e impuestos"),
    ("06", "Bodegas, exhibición e inventario"),
    ("07", "Compras y cuentas por pagar"),
    ("08", "Caja, ventas y comprobantes"),
    ("09", "Devoluciones y notas crédito"),
    ("10", "Cartera y cuentas por cobrar"),
    ("11", "Facturación electrónica con Factus"),
    ("12", "Contabilidad, conciliación y auditoría"),
    ("13", "Reportes, seguridad y copias de respaldo"),
    ("14", "Rutina diaria y preparación para producción"),
]
table = doc.add_table(rows=0, cols=2)
set_table_geometry(table, [900, 8460])
for idx, (num, title) in enumerate(contents):
    row = table.add_row()
    set_cell_text(row.cells[0], num, bold=True, color=CYAN if idx % 2 == 0 else VIOLET, size=10)
    set_cell_text(row.cells[1], title, bold=True, color=NAVY, size=10)
    if idx % 2 == 1:
        for cell in row.cells:
            set_cell_shading(cell, PALE)
set_table_geometry(table, [900, 8460])

add_page_break()

# 1
add_kicker("01 · Visión general")
add_title("Qué es Nubixor")
add_body(
    "Nubixor es una plataforma ERP modular y multiempresa. Reúne administración, catálogo, inventario, compras, caja, cartera, facturación, contabilidad, auditoría y reportes en una sola aplicación."
)
add_callout(
    "Idea central",
    "Cada operación pertenece a una empresa, sucursal, usuario y documento identificables. El sistema no debe mezclar productos, dinero, inventario ni numeración entre empresas.",
    "info",
)
doc.add_heading("Cómo está organizado", level=2)
add_table(
    ["Área", "Propósito", "Resultado principal"],
    [
        ("Administración", "Empresas, sucursales y contexto activo.", "Estructura legal y operativa separada."),
        ("Inventario", "Catálogo, bodegas, exhibición y movimientos.", "Existencias y costo trazables."),
        ("Abastecimiento", "Proveedores, compras y obligaciones.", "Recepción y cuentas por pagar."),
        ("Ventas", "Caja, clientes, pagos, devoluciones y cartera.", "Ventas y documentos por empresa."),
        ("Control", "Usuarios, permisos, contabilidad y auditoría.", "Responsabilidad y evidencia."),
        ("Análisis", "Dashboard y reportes exportables.", "Decisiones con datos operativos."),
    ],
    [1600, 4100, 3660],
)
doc.add_heading("Arquitectura operativa", level=2)
add_flow(["Usuario", "Empresa", "Sucursal", "Operación", "Auditoría"])
add_bullets([
    "La interfaz se ejecuta en el navegador y se conecta a una API local.",
    "PostgreSQL conserva las operaciones y relaciones empresariales.",
    "Redis apoya sesiones, controles temporales y disponibilidad.",
    "Los documentos e imágenes privadas se almacenan fuera de la carpeta pública.",
    "Las migraciones permiten evolucionar la base de datos de manera controlada.",
])

add_page_break()

# 2
add_kicker("02 · Seguridad y personas")
add_title("Acceso, usuarios y roles")
add_body("Nubixor utiliza sesiones seguras, permisos por operación y membresías por empresa. La interfaz oculta áreas no autorizadas, pero la validación definitiva siempre ocurre también en el servidor.")
doc.add_heading("Inicio y protección de sesión", level=2)
add_bullets([
    "La primera apertura permite definir la contraseña inicial del propietario una sola vez.",
    "Las sesiones usan cookie HttpOnly y protección CSRF para operaciones de escritura.",
    "Cinco intentos fallidos producen un bloqueo temporal.",
    "Las invitaciones son personales, de un solo uso y tienen vigencia limitada.",
    "La recuperación de contraseña invalida las sesiones anteriores.",
])
doc.add_heading("Roles recomendados", level=2)
add_table(
    ["Rol", "Uso habitual", "Alcance"],
    [
        ("Propietario", "Configuración completa y decisiones críticas.", "Empresa completa."),
        ("Administrador", "Operación general, usuarios y módulos.", "Empresa o sucursales asignadas."),
        ("Cajero", "Apertura, venta, cobro, comprobantes y cierre.", "Solo Caja & POS."),
        ("Bodega / Operaciones", "Inventario, transferencias y recepción.", "Sucursales asignadas."),
        ("Compras", "Proveedores, órdenes y recepciones.", "Empresa o sucursales."),
        ("Cartera / Contabilidad", "Cobros, pagos, conciliación y revisión.", "Empresa completa según permiso."),
        ("Auditor", "Consulta de evidencia sin alterar operaciones.", "Empresa completa."),
    ],
    [1750, 4580, 3030],
)
add_callout(
    "Cajero",
    "Un usuario de caja ve únicamente el punto de venta y los datos mínimos necesarios para cobrar. Los usuarios con permisos superiores pueden consultar inventario y otras áreas.",
    "success",
)
doc.add_heading("Aislamiento multiempresa", level=2)
add_bullets([
    "Productos, impuestos, clientes, cuentas bancarias y resoluciones pertenecen a una empresa.",
    "Una caja compartida puede mostrar catálogos autorizados de varias empresas.",
    "Un cobro agrupado se divide en ventas y documentos separados por empresa vendedora.",
    "Las transferencias registran la cuenta bancaria receptora de la empresa correcta.",
    "La auditoría conserva quién hizo la operación, cuándo, desde qué empresa y con qué resultado.",
])

# 3-4
add_kicker("03 · Puesta en marcha")
add_title("Configurar una empresa")
add_body("El alta guiada crea una base operativa para empezar: empresa, sucursal principal, bodega, caja e impuestos iniciales. Después deben revisarse los datos legales y operativos.")
add_steps([
    "Crear la empresa con razón social, identificación y datos de contacto.",
    "Revisar el perfil tributario: tipo de documento, responsabilidades, impuestos y tipo de comprobante.",
    "Crear o ajustar sucursales, direcciones y responsables.",
    "Definir bodegas de abastecimiento y ubicaciones de exhibición.",
    "Configurar caja, usuarios autorizados y cuentas bancarias.",
    "Crear categorías, marcas, impuestos y productos.",
    "Cargar existencias mediante recepción, transferencia o ajuste inicial justificado.",
    "Si factura electrónicamente, configurar Factus primero en ambiente TEST.",
])
add_callout(
    "Crative",
    "El escenario actual de Crative usa comprobante interno y productos al 0 %. Su tratamiento tributario debe ser confirmado por el contador; ser persona natural no determina por sí solo todas las obligaciones fiscales.",
    "warning",
)

doc.add_heading("Dashboard ejecutivo", level=1)
add_body("El tablero resume el estado de la empresa activa. Sirve para decidir, no para reemplazar los módulos de detalle.")
add_table(
    ["Indicador", "Qué representa", "Dónde se revisa"],
    [
        ("Ventas del día y mes", "Ingresos y transacciones confirmadas.", "Caja y Reportes."),
        ("Ticket y margen", "Promedio por venta y utilidad bruta estimada.", "Reportes."),
        ("Por cobrar", "Facturas pendientes y vencidas.", "Cartera."),
        ("Por pagar", "Obligaciones próximas y vencidas.", "Cuentas por pagar."),
        ("Valor de inventario", "Existencias valorizadas al costo.", "Inventario."),
        ("Stock bajo", "Productos bajo el mínimo configurado.", "Reposición."),
        ("Flujo proyectado", "Caja + cobros estimados − pagos estimados.", "Dashboard financiero."),
    ],
    [1850, 4550, 2960],
)
add_callout("Lectura correcta", "Una alerta del dashboard conduce al módulo donde se encuentra la evidencia y la acción correspondiente.", "info")

# 5
add_kicker("05 · Catálogo")
add_title("Productos, precios e impuestos")
add_body("El producto pertenece al catálogo de una empresa. Allí viven su identificación, clasificación, costo, precio, tratamiento tributario y fotografía.")
doc.add_heading("Configuración de un producto", level=2)
add_steps([
    "Seleccionar la empresa propietaria del catálogo.",
    "Crear previamente la categoría y la marca cuando apliquen.",
    "Registrar nombre, SKU o código, costo y precio de venta.",
    "Asignar el impuesto validado para esa empresa.",
    "Adjuntar una imagen JPG, PNG o WEBP de hasta 2 MB.",
    "Definir mínimos, máximos y disponibilidad en las ubicaciones correspondientes.",
])
add_callout(
    "Existencia",
    "Crear un producto no crea unidades. Las cantidades se incorporan mediante compras recibidas, transferencias o un ajuste inicial con motivo.",
    "warning",
)
doc.add_heading("Impuestos", level=2)
add_bullets([
    "El tratamiento tributario vive en el producto, no en la bodega.",
    "El historial de una venta conserva el impuesto aplicado en el momento de vender.",
    "Nubixor soporta productos excluidos o al 0 %, pero la clasificación debe validarse por empresa.",
    "Los códigos enviados a Factus provienen de equivalencias configuradas, nunca de ejemplos fijos.",
])
doc.add_heading("Pendientes especializados", level=2)
add_table(
    ["Capacidad", "Estado"],
    [
        ("Unidades de compra, almacenamiento y venta", "Pendiente de implementación completa."),
        ("Variantes con SKU propio", "Pendiente de flujo integral."),
        ("Lotes y vencimientos", "Pendiente."),
        ("Números de serie y garantías", "Pendiente."),
        ("Etiquetas y códigos de barras imprimibles", "Pendiente."),
    ],
    [6100, 3260],
)

# 6
add_kicker("06 · Inventario")
add_title("Bodega, exhibición y trazabilidad")
add_body("Nubixor separa el inventario almacenado del inventario disponible en exhibición. La caja descuenta desde una ubicación de exhibición autorizada, mientras la bodega conserva el respaldo.")
add_flow(["Entrada", "Bodega", "Traslado", "Exhibición", "Venta"])
doc.add_heading("Funciones disponibles", level=2)
add_bullets([
    "Existencias consolidadas por empresa, sucursal, bodega y producto.",
    "Valor al costo y disponibilidad operativa.",
    "Kardex reciente con entradas, salidas, responsables y referencias.",
    "Transferencias transaccionales entre ubicaciones.",
    "Reposición de exhibición según mínimo y máximo.",
    "Ajustes manuales con motivo obligatorio.",
    "Averías, cuarentena, devoluciones y otras novedades.",
    "Órdenes de transferencia pendientes de aceptación.",
])
doc.add_heading("Conteo físico", level=2)
add_steps([
    "Programar una jornada de conteo para una bodega.",
    "Iniciar el conteo y registrar cantidades físicas.",
    "Enviar la toma para revisión.",
    "Completar el conteo; el sistema registra diferencias y movimientos de ajuste.",
])
add_callout(
    "Regla de control",
    "El conteo físico aparece únicamente cuando se programa una toma. No debe utilizarse como formulario cotidiano para aumentar o disminuir existencias.",
    "success",
)

# 7
add_kicker("07 · Abastecimiento")
add_title("Compras y cuentas por pagar")
add_body("El flujo de abastecimiento conecta proveedor, orden, recepción, inventario, costo y obligación.")
add_flow(["Proveedor", "Orden", "Recepción", "Inventario", "Obligación"])
doc.add_heading("Compras", level=2)
add_bullets([
    "Directorio de proveedores con datos y condiciones de pago.",
    "Órdenes con productos, cantidades, costos, impuestos y fechas.",
    "Recepciones parciales o totales.",
    "Actualización transaccional de existencias, kardex y costo promedio.",
    "Seguimiento de unidades pendientes de recibir.",
])
doc.add_heading("Cuentas por pagar", level=2)
add_bullets([
    "Obligaciones originadas en compras recibidas o registradas manualmente.",
    "Fechas de emisión y vencimiento.",
    "Pagos parciales o totales con historial.",
    "Clasificación por edades y proveedor.",
    "Soporte para comparación entre orden, recepción y factura.",
])
add_callout(
    "Buena práctica",
    "No registrar una factura del proveedor como pagada hasta comprobar el medio, la fecha, la referencia y la cuenta bancaria utilizada.",
    "info",
)

add_page_break()

# 8
add_kicker("08 · Caja y ventas")
add_title("Punto de venta y cierre de turno")
add_body("Caja & POS concentra la operación del cajero: abrir turno, seleccionar productos, identificar cliente, cobrar, imprimir y cerrar.")
doc.add_heading("Flujo de una venta", level=2)
add_steps([
    "Abrir el turno con caja, sucursal y fondo inicial contado.",
    "Buscar o escanear productos disponibles en exhibición.",
    "Seleccionar consumidor final o crear/elegir un cliente.",
    "Confirmar cantidades, precios, descuentos autorizados e impuestos.",
    "Elegir contado o crédito.",
    "Registrar efectivo, tarjeta, transferencia o pago mixto.",
    "Confirmar la venta; Nubixor descuenta inventario y crea documentos separados por empresa.",
    "Imprimir o consultar el comprobante desde el historial.",
])
doc.add_heading("Control de caja", level=2)
add_table(
    ["Movimiento", "Efecto", "Evidencia"],
    [
        ("Venta en efectivo", "Aumenta efectivo esperado.", "Venta y comprobante."),
        ("Ingreso manual", "Aumenta caja sin ser venta.", "Motivo y responsable."),
        ("Gasto menor", "Disminuye caja.", "Concepto y valor."),
        ("Retiro", "Disminuye efectivo disponible.", "Motivo y responsable."),
        ("Cierre", "Compara contado contra esperado.", "Arqueo y diferencia."),
    ],
    [2050, 3690, 3620],
)
doc.add_heading("Arqueo y cierre", level=2)
add_bullets([
    "Conteo por denominaciones.",
    "Separación entre efectivo, tarjeta y transferencia.",
    "Cálculo de diferencia contra efectivo esperado.",
    "Justificación obligatoria cuando existe diferencia.",
    "Comprobante e historial de turnos.",
])
add_callout(
    "Venta multiempresa",
    "La caja puede mostrar productos de varias empresas autorizadas. Nubixor realiza el cobro operativo y mantiene ventas, inventarios, comprobantes, impuestos y pagos separados por empresa.",
    "violet",
)

# 9-10
add_kicker("09 · Correcciones")
add_title("Devoluciones y notas crédito")
add_body("La devolución no modifica la venta original. Crea una operación inversa trazable que devuelve dinero, reintegra inventario y conserva la relación fiscal.")
add_flow(["Buscar venta", "Elegir líneas", "Reembolsar", "Reintegrar", "Auditar"])
add_bullets([
    "Permite devoluciones parciales sucesivas o devolución total.",
    "Impide devolver más unidades de las vendidas.",
    "El inventario regresa a la ubicación histórica de salida.",
    "El efectivo genera una salida de caja.",
    "Tarjeta y transferencia exigen referencia; transferencia exige cuenta de la misma empresa.",
    "Una venta electrónica crea una nota crédito pendiente con causal validada.",
    "La operación deja asiento contable y evento de auditoría.",
])
add_callout(
    "Limitación vigente",
    "Las devoluciones de ventas a crédito todavía requieren completar el ajuste automático sobre Cuentas por cobrar. El sistema las bloquea para evitar corromper la cartera.",
    "risk",
)

doc.add_heading("Cartera y cuentas por cobrar", level=1)
add_body("Las ventas a crédito generan una factura en cartera con cliente, vencimiento, saldo e historial.")
add_bullets([
    "Clientes y consumidor final dentro del POS.",
    "Facturas manuales o originadas en ventas a crédito.",
    "Abonos parciales y pagos completos.",
    "Saldo pendiente y documentos vencidos o próximos a vencer.",
    "Consulta de detalle e historial de pagos.",
    "Resumen por edades para seguimiento de cobranza.",
])
add_callout(
    "Control",
    "Un pago de cartera debe indicar valor, fecha, medio y referencia. Si ingresa por transferencia, debe asociarse a la cuenta bancaria receptora correcta.",
    "info",
)

# 11
add_kicker("11 · Facturación electrónica")
add_title("Integración con Factus API V2")
add_body("Nubixor incorpora un conector desacoplado por empresa. Cada empresa posee su propia cuenta, ambiente, credenciales, rangos, resoluciones, equivalencias y documentos.")
doc.add_heading("Capacidades incorporadas", level=2)
add_bullets([
    "Credenciales cifradas y ambiente TEST o PRODUCTION separado por empresa.",
    "OAuth, renovación de token y control preventivo de solicitudes.",
    "Consulta de rangos reales asociados a la cuenta Factus.",
    "Cola idempotente para evitar duplicados.",
    "Creación y validación de facturas.",
    "Soporte para contado, crédito, pagos mixtos e impuesto excluido.",
    "Registro de número, CUFE, QR, respuesta y errores.",
    "Descarga y almacenamiento privado de PDF y XML con hash SHA-256.",
    "Adaptadores para notas crédito y débito.",
    "Reintentos para respuestas transitorias y trazabilidad de intentos.",
])
doc.add_heading("Preparar una empresa", level=2)
add_steps([
    "Configurar la clave de cifrado del servidor.",
    "Registrar credenciales reales de la empresa en ambiente TEST.",
    "Probar la conexión.",
    "Consultar y seleccionar un rango real de la cuenta.",
    "Validar equivalencias DIAN de municipios, unidades, impuestos y medios de pago.",
    "Completar los datos electrónicos de clientes y productos.",
    "Ejecutar pruebas controladas de aceptación, rechazo, crédito y nota crédito.",
    "Solicitar aprobación del contador antes de activar PRODUCTION.",
])
add_callout(
    "Estado real",
    "El conector existe, pero la facturación legal no queda habilitada hasta registrar credenciales reales, resolución vigente, rangos, equivalencias y aprobación tributaria de cada empresa.",
    "warning",
)

# 12
add_kicker("12 · Control financiero")
add_title("Contabilidad, conciliación y auditoría")
add_body("Las operaciones relevantes generan evidencia contable y de auditoría. El objetivo es poder explicar qué ocurrió sin alterar documentos históricos.")
doc.add_heading("Contabilidad disponible", level=2)
add_bullets([
    "Asientos automáticos para ventas, recaudos, compras, pagos y devoluciones habilitadas.",
    "Libro contable y comprobantes imprimibles.",
    "Balance de prueba y auxiliares por cuenta.",
    "Conciliación bancaria con importación y asociación de movimientos.",
    "Períodos contables y bloqueo definitivo.",
    "Expediente mensual para contador y auditor.",
])
doc.add_heading("Auditoría", level=2)
add_table(
    ["Evidencia", "Contenido"],
    [
        ("Evento", "Usuario, empresa, módulo, acción, fecha y resultado."),
        ("Cambio", "Valores anteriores y posteriores cuando aplica."),
        ("Documento", "Referencia a venta, pago, movimiento o asiento."),
        ("Integridad", "Sellos y hashes para detectar alteraciones."),
        ("Exportación", "CSV, evidencia JSON y paquete mensual."),
        ("Revisión", "Registro de validaciones del contador."),
    ],
    [2200, 7160],
)
add_callout(
    "Principio",
    "Las ventas, facturas y asientos no se borran para corregir errores. Se crean devoluciones, notas, reversos o ajustes enlazados.",
    "success",
)

add_page_break()

# 13
add_kicker("13 · Información y continuidad")
add_title("Reportes, archivos y copias")
doc.add_heading("Reportes operativos", level=2)
add_bullets([
    "Ventas por producto, categoría, caja, sucursal y medio de pago.",
    "Rentabilidad y margen bruto.",
    "Inventario, valorización y rotación.",
    "Compras y desempeño por proveedor.",
    "Cartera y cuentas por pagar por edades.",
    "Filtros por fecha y sucursal, búsqueda, paginación y exportación CSV.",
])
doc.add_heading("Archivos protegidos", level=2)
add_bullets([
    "Imágenes y documentos se guardan con nombres aleatorios fuera del directorio público.",
    "Cada archivo conserva empresa, responsable, fecha, tipo y hash SHA-256.",
    "El acceso exige una sesión con membresía autorizada.",
    "PDF, XML y soportes fiscales deben formar parte de la copia de seguridad.",
])
doc.add_heading("Copias de seguridad", level=2)
add_bullets([
    "La copia automática incluye PostgreSQL y almacenamiento privado.",
    "Se cifra con AES-256-GCM y genera un hash SHA-256.",
    "La retención predeterminada es de 30 días.",
    "Una copia no se considera confiable hasta ensayar su restauración en una base separada.",
])
add_callout(
    "Producción",
    "Las copias deben enviarse también a un destino externo. Una copia almacenada únicamente en el mismo servidor no protege contra pérdida total del equipo.",
    "warning",
)

add_page_break()

# 14
add_kicker("14 · Operación")
add_title("Rutina recomendada del negocio")
doc.add_heading("Inicio del día", level=2)
add_steps([
    "Comprobar que Nubixor, PostgreSQL y Redis estén disponibles.",
    "Revisar alertas del dashboard.",
    "Confirmar empresa, sucursal, usuario y caja.",
    "Abrir turno contando el fondo inicial.",
    "Revisar reposición de exhibición y productos sin disponibilidad.",
])
doc.add_heading("Durante la operación", level=2)
add_bullets([
    "Registrar cada venta con su medio de pago real.",
    "Exigir referencia y cuenta receptora en transferencias.",
    "Registrar ingresos, gastos y retiros cuando ocurren.",
    "Recibir compras contra su orden y documento.",
    "Usar devoluciones o notas para corregir, sin editar la operación original.",
])
doc.add_heading("Cierre del día", level=2)
add_steps([
    "Contar el efectivo por denominaciones.",
    "Explicar cualquier diferencia.",
    "Cerrar el turno y conservar el comprobante.",
    "Comparar transferencias y tarjetas contra sus cuentas.",
    "Revisar ventas, inventario, cartera, obligaciones y documentos rechazados.",
    "Confirmar que la copia de seguridad fue generada.",
])

doc.add_heading("Qué falta antes de producción", level=1)
add_table(
    ["Prioridad", "Trabajo pendiente", "Condición de cierre"],
    [
        ("1", "Devolución de ventas a crédito.", "Ajuste correcto de cartera y nota crédito."),
        ("2", "Factus real por empresa.", "Credenciales, rango, resolución y equivalencias."),
        ("3", "Pruebas integrales.", "Venta, concurrencia, cierre, conciliación y restauración."),
        ("4", "Inventario especializado.", "Unidades, variantes, lotes, series y garantías."),
        ("5", "Infraestructura.", "Dominio, HTTPS, correo y copia externa."),
        ("6", "Validación tributaria.", "Aprobación independiente del contador."),
    ],
    [1050, 4040, 4270],
)
add_callout(
    "Conclusión",
    "Nubixor ya puede demostrarse y utilizarse en pruebas controladas. Para operar con dinero y facturación legal debe completarse la habilitación por empresa, las pruebas de extremo a extremo y la aprobación contable.",
    "violet",
)

# Cierre y glosario
add_kicker("Referencia rápida")
add_title("Glosario y soporte")
glossary_table = add_table(
    ["Término", "Significado"],
    [
        ("Empresa activa", "Contexto legal desde el cual se consultan y registran datos."),
        ("Sucursal", "Punto operativo perteneciente a una empresa."),
        ("Bodega", "Ubicación física de almacenamiento."),
        ("Exhibición", "Ubicación desde la cual Caja puede vender."),
        ("Kardex", "Historial cronológico de entradas y salidas de inventario."),
        ("Turno", "Período de responsabilidad de un cajero sobre una caja."),
        ("Arqueo", "Comparación entre dinero contado y dinero esperado."),
        ("CUFE", "Código Único de Factura Electrónica asignado al documento."),
        ("Nota crédito", "Documento que disminuye o reversa una factura."),
        ("Idempotencia", "Protección que evita repetir una operación al reenviarla."),
        ("Conciliación", "Comparación entre movimientos del sistema y del banco."),
        ("Expediente mensual", "Paquete de evidencia para contador o auditor."),
    ],
    [2200, 7160],
    font_size=8.6,
)
for glossary_row in glossary_table.rows:
    for glossary_cell in glossary_row.cells:
        set_cell_margins(glossary_cell, 55, 120, 55, 120)
doc.add_heading("Comprobaciones técnicas", level=2)
for technical_item in [
    "Aplicación local: http://localhost:4100",
    "Salud: /api/health · Dependencias: /api/health/ready",
    "Descripción y cierre operativo: README.md · docs/PLAN_CIERRE_OPERATIVO.md",
    "Producción y Factus: docs/PUESTA_EN_PRODUCCION.md · docs/FACTUS_V2.md",
]:
    p = doc.add_paragraph()
    set_num(p, bullet_num_id)
    set_paragraph_spacing(p, after=1, line=1.1)
    run = p.add_run(technical_item)
    set_run_font(run, BODY_FONT, 9, INK)
# Metadatos y guardado
doc.core_properties.title = "Manual funcional y operativo de Nubixor"
doc.core_properties.subject = "Funcionamiento, módulos, controles y preparación operativa"
doc.core_properties.author = "Nubixor"
doc.core_properties.keywords = "Nubixor, ERP, POS, inventario, facturación, auditoría"
doc.core_properties.comments = "Documento generado a partir del estado funcional del proyecto Nubixor."

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
